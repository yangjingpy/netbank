from docx import Document
import os
import re
import csv

def search_title(file_path):
    # CSV test plan template
    csv_template_title = ['用例编号', '用例描述', '测试步骤', '期望结果', '优先级', '影响版本']
    test_case_version = '4.0'
    test_case_pri_map = {
        'P1': 'P1',
        'P2': 'P2',
        'P3': 'P3',
        'P4': 'P4'
    }

    #open docx
    document = Document(file_path)
    csv_name = os.path.join(os.path.dirname(FILE_PATH),docx_name + '.csv')

    paras = document.paragraphs
    # newline否则每行默认输入一行空格
    f = open(csv_name, 'w', encoding='utf-8', newline='')
    writer = csv.writer(f)

    #read tables
    tables = document.tables
    print(len(tables))
    #table = tables[1]
    n = 0
    #从第二个表格开始读起
    for table in tables[1:]:
        table_content = dict()
        for i in range(len(table.rows)):
            #There are four columns in the second row
            if i == 1:
                # # cell(i,0)表示第(i+1)行第1列数据，以此类推
                table_content[table.cell(i, 0).text] = table.cell(i, 1).text
                table_content[table.cell(i, 2).text] = table.cell(i, 3).text
            table_content[table.cell(i,0).text] = table.cell(i,1).text
        if n == 0:
            list_keys = [i for i in table_content.keys()]
            writer.writerow(csv_template_title)
        
        test_case_item = list()
        test_case_item.append(table_content['用例编号'])
        test_case_des = '[描述]\n' + table_content['用例描述'] + '\n' + \
                        '[预置条件]\n' + table_content['预置条件'] + '\n' + \
                        '[拓扑描述]\n' + table_content['使用的网络图'] + '\n' + \
                        '[支持平台]\nALL\n[其他]'
        test_case_item.append(test_case_des)
        test_case_item.append(table_content['测试步骤'])
        test_case_item.append(table_content['期望结果'])
        test_case_item.append(test_case_pri_map[table_content['优先级']])
        test_case_item.append(test_case_version)
        
        writer.writerow(test_case_item)
        n+=1
    f.close()

docx_testplan_list = [r'D:\用例\网管web4.0.docx']
for tp in docx_testplan_list:
    FILE_PATH = os.path.abspath(tp)
    ## get the name of docx
    res1 = re.match('(.+)\.docx', os.path.basename(FILE_PATH))
    ## group() is all,group(1) is first
    docx_name = res1.group(1)
    doc_result = search_title(FILE_PATH)

